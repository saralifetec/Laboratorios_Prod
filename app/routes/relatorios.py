from copy import deepcopy
from pathlib import Path
import re
from PIL import Image
import tempfile
import unicodedata
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from flask import Blueprint, current_app, jsonify, request, send_file
from docx.shared import Cm
import win32com
from app.models import Ensaio, Testes
from docx.enum.table import WD_TABLE_ALIGNMENT
from services.relatorios_utils import (
    docx_to_pdf,
    formatar_celula,
    ler_resistencias,
    obter_numero_peca,
    prepare_image_for_report,
    resolver_base_path,
    obter_pecas,
    extract_peca_label,
    get_before_after_images,
    find_child_dir
)

relatorios_bp = Blueprint('relatorios', __name__)


@relatorios_bp.route('/relatorios/gerar', methods=['POST'])
def gerar_relatorio():

    data = request.get_json()

    tipo = data.get("tipo")
    ensaio = data.get("ensaio")
    testes = data.get("testes", [])

    if not ensaio:
        return jsonify({"error_key": "msg.preencha_ensaio"}), 400

    try:

        if tipo == "ensaio":
            #  usa o endpoint já existente
            return relatorio_fotos()

        elif tipo == "ensaio_teste":
            return relatorio_fotos_por_teste()

        elif tipo == "setup":
            return relatorio_setup_por_teste()
        
        elif tipo == "resistencias":
            return relatorio_resistencias()

        else:
            return jsonify({"error_key": "msg.tipo_relatorio_invalido"}), 400

    except Exception as e:
       

        return jsonify({
            "error_key": "msg.erro_gerar_relatorio",
            "detail": str(e)
        }), 500


# =================================================
# VERIFICAR SE JÁ EXISTE
# =================================================
@relatorios_bp.route('/relatorios/existe', methods=['POST'])
def verificar_relatorio_existente():

    data = request.get_json()
    ensaio_numero = data.get("ensaio")
    tipo = data.get("tipo")

    if not ensaio_numero or not tipo:
        return jsonify({"error": "Parâmetros inválidos"}), 400

    #  OBTER O ENSAIO (faltava isto!)
    ensaio = Ensaio.query.filter_by(ensaio=ensaio_numero).first()

    if not ensaio:
        return jsonify({"error": "Ensaio não encontrado"}), 404

    #  usar a instância correta
    base_path = resolver_base_path(ensaio)

    if not base_path:
        return jsonify({"error": "Pasta do ensaio não encontrada"}), 400

    #  pasta correta
    output_dir = base_path / "Reports_Temp"
    output_dir.mkdir(parents=True, exist_ok=True)

    #  nome do ficheiro
    if tipo == "ensaio":
        filename = f"Relatorio_Fotos_{ensaio_numero}.docx"
    elif tipo == "ensaio_teste":

        testes_sel = [int(t) for t in data.get("testes", [])]

        testes = Testes.query.filter_by(ensaio_id=ensaio.id).all()

        testes_com_pasta = [
            t for t in sorted(testes, key=lambda x: x.ordem or 0)
            if t.teste and t.teste.criarpasta
        ]

        testes_validos = [
            t for t in testes_com_pasta
            if not testes_sel or t.ordem in testes_sel
        ]

        testes_existentes = []

        for teste in testes_validos:

            numero_pasta = testes_com_pasta.index(teste) + 1
            nome_clean = teste.teste.teste.upper().replace(" ", "_")

            filename = f"{ensaio_numero}_{numero_pasta}_{nome_clean}.docx"
            output_path = output_dir / filename

            if output_path.exists():
                testes_existentes.append(
                    f"{numero_pasta} - {teste.teste.teste}"
                )

        return jsonify({
            "exists": len(testes_existentes) > 0,
            "testes_existentes": testes_existentes
        })
        
    elif tipo == "setup":

        testes_sel = [int(t) for t in data.get("testes", [])]

        testes = Testes.query.filter_by(ensaio_id=ensaio.id).all()

        testes_com_pasta = [
            t for t in sorted(testes, key=lambda x: x.ordem or 0)
            if t.teste and t.teste.criarpasta
        ]

        testes_validos = [
            t for t in testes_com_pasta
            if not testes_sel or t.ordem in testes_sel
        ]

        testes_existentes = []

        for teste in testes_validos:

            numero_pasta = testes_com_pasta.index(teste) + 1
            nome_clean = teste.teste.teste.upper().replace(" ", "_")

            filename = f"{ensaio_numero}_{numero_pasta}_{nome_clean}_Setup.docx"
            output_path = output_dir / filename

            if output_path.exists():
                testes_existentes.append(
                    f"{numero_pasta} - {teste.teste.teste}"
                )

        return jsonify({
            "exists": len(testes_existentes) > 0,
            "testes_existentes": testes_existentes
        })
    else:
        filename = f"Relatorios_{ensaio_numero}.docx"

    output_path = output_dir / filename

    return jsonify({
        "exists": output_path.exists()
    })

# =================================================
# RELATÓRIO FOTOS COMPLETO
# =================================================
@relatorios_bp.route('/relatorio_fotos', methods=['POST'])
def relatorio_fotos():

    # =================================================
    # TEMPLATE
    # =================================================
    TEMPLATE_PATH = Path(
        r"W:\DEPARTMENTS\LAB\PROCESS\01 QUALITY\01.01 DOCUMENTATION"
        r"\01.01.05 TEMPLATES\ALL\REPORTS\GBS\Reports_Fotos_Ensaios.docx"
    )

    if not TEMPLATE_PATH.exists():
        return jsonify({"error_key": "msg.template_reports_fotos_nao_encontrado"}), 500

    # =================================================
    # INPUT
    # =================================================
    data = request.get_json()
    ensaio_numero = data.get("ensaio")

    if not ensaio_numero:
        return jsonify({"error_key": "msg.preencha_ensaio"}), 400

    ensaio = Ensaio.query.filter_by(ensaio=ensaio_numero).first()
    if not ensaio:
        return jsonify({"error_key": "msg.ensaio_nao_encontrado"}), 404

    # =================================================
    # BASE PATH (utils)
    # =================================================
    base_path = resolver_base_path(ensaio)

    if not base_path:
        return jsonify({"error_key": "msg.pasta_nao_encontrada"}), 400

    before_test_base = base_path / "BEFORE TEST"

    if not before_test_base.exists():
        return jsonify({"error_key": "msg.pasta_before_test_nao_existe"}), 400

    pecas = obter_pecas(before_test_base)

    if not pecas:
        return jsonify({"error_key": "msg.nao_existem_pecas"}), 400

    # =================================================
    # TESTES
    # =================================================
    testes = Testes.query.filter_by(ensaio_id=ensaio.id).all()

    testes_validos = [
        t for t in sorted(testes, key=lambda x: x.ordem or 0)
        if t.teste and t.teste.criarpasta
    ]

    # =================================================
    # DOCUMENTO
    # =================================================
    doc = Document(TEMPLATE_PATH)

    if not doc.tables:
        return jsonify({"error": "Template inválido (sem tabela)"}), 500

    template_table = doc.tables[0]

    # limpar parágrafos iniciais
    while len(doc.paragraphs) > 0:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    last_paragraph = doc.add_paragraph()
    previous_after_test_path = None

    # =================================================
    # LOOP TESTES
    # =================================================
    for idx, teste in enumerate(testes_validos):

        # --- TÍTULO
        titulo = doc.add_paragraph(teste.teste.teste.upper())
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_paragraph = doc.add_paragraph("")

        numero_pasta = idx + 1
        nome_pasta = f"{numero_pasta} - {teste.teste.teste.upper()}"

        pasta_teste = find_child_dir(base_path, nome_pasta)

        if not pasta_teste:
            continue

        teste_after = (
            pasta_teste / "AFTER TEST"
            if (pasta_teste / "AFTER TEST").exists()
            else None
        )

        # =================================================
        # LOOP PEÇAS
        # =================================================
        for peca in pecas:

            label = extract_peca_label(peca.name)

            before_imgs, after_imgs = get_before_after_images(
                peca,
                before_test_base,
                previous_after_test_path,
                teste_after
            )

            new_tbl_xml = deepcopy(template_table._tbl)
            last_paragraph._p.addnext(new_tbl_xml)

            tbl = Table(new_tbl_xml, doc)

            # headers
            p_left = tbl.rows[0].cells[0].paragraphs[0]
            p_left.text = f"{label} Before Test"
            p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p_right = tbl.rows[0].cells[1].paragraphs[0]
            p_right.text = f"{label} After Test"
            p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # imagens
            for i in range(3):

                if i < len(before_imgs):
                    try:
                        img = prepare_image_for_report(before_imgs[i])
                        tbl.rows[i+1].cells[0].paragraphs[0].add_run() \
                            .add_picture(str(img), width=Cm(8))
                    except Exception as e:
                        print("Erro imagem BEFORE:", e)

                if i < len(after_imgs):
                    try:
                        img = prepare_image_for_report(after_imgs[i])
                        tbl.rows[i+1].cells[1].paragraphs[0].add_run() \
                            .add_picture(str(img), width=Cm(8))
                    except Exception as e:
                        print("Erro imagem AFTER:", e)

            # quebra de página
            doc.add_page_break()
            last_paragraph = doc.paragraphs[-1]

        previous_after_test_path = teste_after

    # =================================================
    # REMOVER TABELA TEMPLATE
    # =================================================
    try:
        doc._body._body.remove(template_table._tbl)
    except Exception:
        pass

    # =================================================
    # GUARDAR
    # =================================================
    
    
    output_dir = base_path / "Reports_Temp"
    output_dir.mkdir(parents=True, exist_ok=True)

    output_path = output_dir / f"Relatorio_Fotos_{ensaio_numero}.docx"

    doc.save(output_path)


    # =================================================
    # DEVOLVER (SEM DOWNLOAD)
    # =================================================
    return jsonify({
        "success": True,
        "message_key": "msg.relatorio_gerado_sucesso",
        "params": {
            "path": str(output_path)
        }
    })


# ============================================
# RELATÓRIO RESISTÊNCIAS
# ============================================
@relatorios_bp.route('/relatorio_resistencias', methods=['POST'])
def relatorio_resistencias():

    try:

        TEMPLATE_PATH = Path(
            r"W:\DEPARTMENTS\LAB\PROCESS\01 QUALITY\01.01 DOCUMENTATION"
            r"\01.01.05 TEMPLATES\ALL\REPORTS\GBS\Reports_Resistencias.docx"
        )

        if not TEMPLATE_PATH.exists():
            return jsonify({"error_key": "msg.template_reports_resistencias_nao_encontrado"}), 500

        data = request.get_json()
        ensaio_numero = data.get("ensaio")

        if not ensaio_numero:
            return jsonify({"error_key": "msg.preencha_ensaio"}), 400

        ensaio = Ensaio.query.filter_by(ensaio=ensaio_numero).first()

        if not ensaio:
            return jsonify({"error_key": "msg.ensaio_nao_encontrado"}), 404

        base_path = resolver_base_path(ensaio)

        if not base_path:
            return jsonify({"error_key": "msg.pasta_nao_encontrada"}), 400

        before_test_base = base_path / "BEFORE TEST"

        if not before_test_base.exists():
            return jsonify({"error_key": "msg.pasta_before_test_nao_existe"}), 400

        pecas = obter_pecas(before_test_base)

        if not pecas:
            return jsonify({"error_key": "msg.nao_existem_pecas"}), 400

        csv_path = base_path / "Resistance before test.csv"

        if not csv_path.exists():
            return jsonify({"error_key": "msg.ficheiro_resistencia_nao_encontrado"}), 400

        resistencias_before = ler_resistencias(csv_path)

        # =================================================
        # TESTES
        # =================================================

        testes = Testes.query.filter_by(ensaio_id=ensaio.id).all()

        


        testes_validos = [
            t for t in sorted(testes, key=lambda x: x.ordem or 0)
            if t.teste and t.teste.criarpasta
        ]

        

        # =================================================
        # DADOS RELATÓRIO
        # =================================================

        dados_relatorio = {}

        for peca in pecas:

            label_peca = extract_peca_label(peca.name)
            numero_csv = str(obter_numero_peca(label_peca))

            dados_relatorio[label_peca] = {
                "before": {},
                "testes": {}
            }


            for ignitor, dados in resistencias_before.items():

                cor = dados["cor"]

                dados_relatorio[label_peca]["before"][cor] = (
                    dados["pecas"].get(numero_csv)
                )

        # =================================================
        # LER TODOS OS TESTES
        # =================================================

        for idx, teste in enumerate(testes_validos):

            numero_pasta = idx + 1

            nome_pasta = f"{numero_pasta} - {teste.teste.teste.upper()}"

            pasta_teste = find_child_dir(base_path, nome_pasta)

            if not pasta_teste:
                continue

            csv_teste = pasta_teste / "Resistance after test.csv"

            if not csv_teste.exists():
                continue

            resistencias_teste = ler_resistencias(csv_teste)

            nome_teste = teste.teste.teste

            for label_peca in dados_relatorio:

                numero_csv = str(
                    obter_numero_peca(label_peca)
                )

                
                dados_relatorio[label_peca]["testes"][nome_teste] = {}

                if not pasta_teste:
                    continue

                csv_teste = pasta_teste / "Resistance after test.csv"

                if not csv_teste.exists():
                    continue


                # caso simples: apenas um ignitor
                if len(resistencias_before) == 1 and len(resistencias_teste) == 1:

                    ign_teste = next(iter(resistencias_teste.values()))

                    valor = ign_teste["pecas"].get(numero_csv)

                    dados_relatorio[label_peca]["testes"][nome_teste][None] = valor

                # vários ignitores -> procurar pela cor
                else:

                    for ign_before, dados_before in resistencias_before.items():

                        cor_before = (
                            dados_before["cor"] or ""
                        ).lower()

                        valor = None

                        for ign_teste, dados_teste in resistencias_teste.items():

                            cor_teste = (
                                dados_teste["cor"] or ""
                            ).lower()

                            if cor_teste == cor_before:

                                valor = dados_teste["pecas"].get(
                                    numero_csv
                                )

                                break

                        dados_relatorio[label_peca]["testes"][
                            nome_teste
                        ][cor_before] = valor



        doc = Document(TEMPLATE_PATH)

        if not doc.tables:
            return jsonify({"error": "Template inválido (sem tabela)"}), 500

       
        # =================================================
        # RESISTÊNCIAS BEFORE TEST
        # =================================================

        ignitores = list(resistencias_before.keys())
        num_ignitores = len(ignitores)

        # =================================================
        # REMOVER TABELA TEMPLATE
        # =================================================

        if doc.tables:

            tabela_template = doc.tables[0]

            tabela_template._element.getparent().remove(
                tabela_template._element
            )

        # =================================================
        # CRIAR NOVA TABELA
        # =================================================

       
        # =================================================
        # COLUNAS
        # =================================================

        ignitores_ordenados = list(
            resistencias_before.keys()
        )

        num_ignitores = len(ignitores_ordenados)

        blocos = ["Before Test"] + [
            t.teste.teste
            for t in testes_validos
        ]

        num_colunas = 1 + (
            len(blocos) * num_ignitores
        )

        

        tabela = doc.add_table(
            rows=2,
            cols=num_colunas
        )

        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

        tabela.style = "Table Grid"
       
        header1 = tabela.rows[0]
        header2 = tabela.rows[1]

        

        # =================================================
        # CABEÇALHO
        # =================================================

        # Samples

        header1.cells[0].merge(header2.cells[0])
        header1.cells[0].text = "Samples"

        coluna = 1

        print("BLOCOS", blocos)
        print("IGNITORES", num_ignitores)
        print("COLUNAS", num_colunas)

        for bloco in blocos:

            print("BLOCO", bloco)

            inicio = coluna
            fim = coluna + num_ignitores - 1

            if num_ignitores > 1:

                header1.cells[inicio].merge(
                    header1.cells[fim]
                )

                header1.cells[inicio].text = bloco

            else:

                header1.cells[inicio].text = bloco

            for ign in range(num_ignitores):

                header2.cells[
                    coluna + ign
                ].text = str(ign + 1)

            coluna += num_ignitores

        for row in [header1, header2]:

            for cell in row.cells:

                formatar_celula(
                    cell,
                    header=True,
                    bold=True,
                    center=True
                )

        # =================================================
        # PREENCHER PEÇAS
        # =================================================

        for label_peca in sorted(
            dados_relatorio.keys(),
            key=obter_numero_peca
        ):
            
            
            row = tabela.add_row()

            
            row.cells[0].text = label_peca

            formatar_celula(
                row.cells[0],
                center=True
            )

            coluna = 1

            # BEFORE TEST

            for ignitor in ignitores_ordenados:

                print("IGNITOR", ignitor)
                cor = resistencias_before[
                    ignitor
                ]["cor"]

                print("COR", cor)

                valor = dados_relatorio[
                    label_peca
                ]["before"].get(cor)


                if valor is not None:

                    row.cells[coluna].text = (
                        f"{valor:.2f}"
                    )

            
                
                formatar_celula(
                    row.cells[coluna],
                    center=True
                )



                coluna += 1

            # TESTES

            for nome_teste in blocos[1:]:

                for ignitor in ignitores_ordenados:

                    cor = resistencias_before[
                        ignitor
                    ]["cor"]

                    valor = (
                        dados_relatorio[
                            label_peca
                        ]["testes"]
                        .get(nome_teste, {})
                        .get(cor)
                    )

                    if valor is not None:

                        row.cells[coluna].text = (
                            f"{valor:.2f}"
                        )

                    formatar_celula(
                        row.cells[coluna],
                        center=True
                    )

                    coluna += 1

        # =================================================
        # GUARDAR
        # =================================================

        output_dir = base_path / "Reports_Temp"
        output_dir.mkdir(parents=True, exist_ok=True)

        docx_path = output_dir / f"Relatorio_Resistencias_{ensaio_numero}.docx"

        print("ANTES SAVE")
        doc.save(docx_path)
        

        pdf_path = output_dir / f"Relatorio_Resistencias_{ensaio_numero}.pdf"

        docx_to_pdf(docx_path, pdf_path)

        return jsonify({
            "success": True,
            "message_key": "msg.relatorio_gerado_sucesso",
            "params": {
                "path": str(pdf_path)
            }
        })
    
    except Exception as e:

        traceback.print_exc()

        print(f"ERRO: {e}")

        raise





# =================================================
# RELATORIO FOTOS POR TESTE
# =================================================
from docx.table import Table

@relatorios_bp.route('/relatorio_fotos_por_teste', methods=['POST'])
def relatorio_fotos_por_teste():

    TEMPLATE_PATH = Path(
        r"W:\DEPARTMENTS\LAB\PROCESS\01 QUALITY\01.01 DOCUMENTATION"
        r"\01.01.05 TEMPLATES\ALL\REPORTS\GBS\Reports_Fotos_Ensaios.docx"
    )

    if not TEMPLATE_PATH.exists():
        return jsonify({"error_key": "msg.template_reports_fotos_nao_encontrado"}), 500

    data = request.get_json()
    ensaio_numero = data.get("ensaio")
    testes_sel = [int(t) for t in data.get("testes", [])]

    if not ensaio_numero:
        return jsonify({"error_key": "msg.preencha_ensaio"}), 400

    ensaio = Ensaio.query.filter_by(ensaio=ensaio_numero).first()
    if not ensaio:
        return jsonify({"error_key": "msg.ensaio_nao_encontrado"}), 404

    # =================================================
    # BASE PATH
    # =================================================
    base_path = resolver_base_path(ensaio)
    if not base_path:
        return jsonify({"error_key": "msg.pasta_nao_encontrada"}), 400

    before_test_base = base_path / "BEFORE TEST"
    if not before_test_base.exists():
        return jsonify({"error_key": "msg.pasta_before_test_nao_existe"}), 400

    pecas = obter_pecas(before_test_base)
    if not pecas:
        return jsonify({"error_key": "msg.nao_existem_pecas"}), 400

    # =================================================
    # TESTES
    # =================================================
    testes = Testes.query.filter_by(ensaio_id=ensaio.id).all()

    testes_com_pasta = [
        t for t in sorted(testes, key=lambda x: x.ordem or 0)
        if t.teste and t.teste.criarpasta
    ]

    testes_validos = [
        t for t in testes_com_pasta
        if not testes_sel or t.ordem in testes_sel
    ]

    if not testes_validos:
        return jsonify({"error_key": "msg.sem_testes_selecionados"}), 400

    # =================================================
    # OUTPUT
    # =================================================
    output_dir = base_path / "Reports_Temp"
    output_dir.mkdir(parents=True, exist_ok=True)

    # =================================================
    # LOOP TESTES
    # =================================================
    for teste in testes_validos:

        # ✅ numeração real
        idx_real = testes_com_pasta.index(teste)
        numero_pasta = idx_real + 1

        doc = Document(TEMPLATE_PATH)

        if not doc.tables:
            return jsonify({"error": "Template inválido (sem tabela)"}), 500

        template_table = doc.tables[0]

        # ✅ limpar template
        while len(doc.paragraphs) > 0:
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)

        last_paragraph = doc.add_paragraph()

        # ---------------------------
        # TÍTULO
        # ---------------------------
        titulo = doc.add_paragraph(teste.teste.teste.upper())
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_paragraph = doc.add_paragraph("")

        # ---------------------------
        # PASTA ATUAL
        # ---------------------------
        nome_pasta = f"{numero_pasta} - {teste.teste.teste.upper()}"
        pasta_teste = find_child_dir(base_path, nome_pasta)

        if not pasta_teste:
            continue

        teste_after = (
            pasta_teste / "AFTER TEST"
            if (pasta_teste / "AFTER TEST").exists()
            else None
        )

        # =================================================
        # ✅ BEFORE CORRETO (BASEADO NA SEQUÊNCIA REAL)
        # =================================================
        if idx_real == 0:
            previous_after_test_path = None
        else:
            teste_anterior = testes_com_pasta[idx_real - 1]

            numero_pasta_anterior = idx_real
            nome_pasta_anterior = f"{numero_pasta_anterior} - {teste_anterior.teste.teste.upper()}"

            pasta_anterior = find_child_dir(base_path, nome_pasta_anterior)

            if pasta_anterior and (pasta_anterior / "AFTER TEST").exists():
                previous_after_test_path = pasta_anterior / "AFTER TEST"
            else:
                previous_after_test_path = None

        # =================================================
        # LOOP PEÇAS
        # =================================================
        for peca in pecas:

            label = extract_peca_label(peca.name)

            before_imgs, after_imgs = get_before_after_images(
                peca,
                before_test_base,
                previous_after_test_path,
                teste_after
            )

            # ✅ tabela correta
            new_tbl_xml = deepcopy(template_table._tbl)
            last_paragraph._p.addnext(new_tbl_xml)
            tbl = Table(new_tbl_xml, doc)

            # HEADERS
            p_left = tbl.rows[0].cells[0].paragraphs[0]
            p_left.text = f"{label} Before Test"
            p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p_right = tbl.rows[0].cells[1].paragraphs[0]
            p_right.text = f"{label} After Test"
            p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # IMAGENS
            for i in range(3):

                if i < len(before_imgs):
                    try:
                        img_before = prepare_image_for_report(before_imgs[i])
                        tbl.rows[i+1].cells[0].paragraphs[0].add_run() \
                            .add_picture(str(img_before), width=Cm(8))
                    except Exception as e:
                        print("Erro BEFORE:", e)

                if i < len(after_imgs):
                    try:
                        img_after = prepare_image_for_report(after_imgs[i])
                        tbl.rows[i+1].cells[1].paragraphs[0].add_run() \
                            .add_picture(str(img_after), width=Cm(8))
                    except Exception as e:
                        print("Erro AFTER:", e)

            doc.add_page_break()
            last_paragraph = doc.paragraphs[-1]

        # remover tabela base
        try:
            doc._body._body.remove(template_table._tbl)
        except Exception:
            pass

        nome_clean = teste.teste.teste.upper().replace(" ", "_")
        output_path = output_dir / f"{ensaio_numero}_{numero_pasta}_{nome_clean}.docx"

        doc.save(output_path)

    # =================================================
    # RESULTADO
    # =================================================
    return jsonify({
        "success": True,
        "message_key": "msg.relatorios_criados_sucesso",
        "params": {"path": str(output_dir)}
    })



# =================================================
# RELATÓRIO SETUP
# =================================================
@relatorios_bp.route('/relatorio_setup_por_teste', methods=['POST'])
def relatorio_setup_por_teste():

    TEMPLATE_PATH = Path(
        r"W:\DEPARTMENTS\LAB\PROCESS\01 QUALITY\01.01 DOCUMENTATION"
        r"\01.01.05 TEMPLATES\ALL\REPORTS\GBS\Reports_Fotos_Setup.docx"
    )

    if not TEMPLATE_PATH.exists():
        return jsonify({"error_key": "msg.template_reports_fotos_nao_encontrado"}), 500

    data = request.get_json()
    ensaio_numero = data.get("ensaio")
    testes_sel = [int(t) for t in data.get("testes", [])]

    if not ensaio_numero:
        return jsonify({"error_key": "msg.preencha_ensaio"}), 400

    ensaio = Ensaio.query.filter_by(ensaio=ensaio_numero).first()
    if not ensaio:
        return jsonify({"error_key": "msg.ensaio_nao_encontrado"}), 404

    # =================================================
    # BASE PATH
    # =================================================
    base_path = resolver_base_path(ensaio)
    if not base_path:
        return jsonify({"error_key": "msg.pasta_nao_encontrada"}), 400


    # =================================================
    # TESTES
    # =================================================
    testes = Testes.query.filter_by(ensaio_id=ensaio.id).all()

    testes_com_pasta = [
        t for t in sorted(testes, key=lambda x: x.ordem or 0)
        if t.teste and t.teste.criarpasta
    ]

    testes_validos = [
        t for t in testes_com_pasta
        if not testes_sel or t.ordem in testes_sel
    ]

    if not testes_validos:
        return jsonify({"error_key": "msg.sem_testes_selecionados"}), 400

    # =================================================
    # OUTPUT
    # =================================================
    output_dir = base_path / "Reports_Temp"
    output_dir.mkdir(parents=True, exist_ok=True)

    # =================================================
    # LOOP TESTES
    # =================================================
    for teste in testes_validos:

        # numeração real
        idx_real = testes_com_pasta.index(teste)
        numero_pasta = idx_real + 1

        doc = Document(TEMPLATE_PATH)

        if not doc.tables:
            return jsonify({"error": "Template inválido (sem tabela)"}), 500

        template_table = doc.tables[0]

        # limpar template
        while len(doc.paragraphs) > 0:
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)

        last_paragraph = doc.add_paragraph()

        # ---------------------------
        # TÍTULO
        # ---------------------------
        titulo = doc.add_paragraph(teste.teste.teste.upper())
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_paragraph = doc.add_paragraph("")

        # ---------------------------
        # PASTA ATUAL
        # ---------------------------
        nome_pasta = f"{numero_pasta} - {teste.teste.teste.upper()}"
        pasta_teste = find_child_dir(base_path, nome_pasta)

        if not pasta_teste:
            continue

        pasta_setup = (
            pasta_teste / "SETUP"
            if (pasta_teste / "SETUP").exists()
            else None
        )

        #==============================
        # OBTER IMAGENS
        #==============================
        setup_images = []

        if pasta_setup:
            setup_images = sorted([
                p for p in pasta_setup.glob("*.*")
                if p.suffix.lower() in [".jpg", ".jpeg", ".png", ".bmp"]
            ])


        # =================================================
        # LOOP IMAGENS
        # =================================================
        # =================================================
        # LOOP IMAGENS
        # =================================================
        for img_path in setup_images:

            #  criar tabela nova
            new_tbl_xml = deepcopy(template_table._tbl)
            last_paragraph._p.addnext(new_tbl_xml)
            tbl = Table(new_tbl_xml, doc)

            # ---------------------
            # TÍTULO (linha 0)
            # ---------------------
            cell_title = tbl.rows[0].cells[0]
            cell_title.text = "Setup"
            cell_title.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # ---------------------
            # IMAGEM (linha 1)
            # ---------------------
            try:
                img_processed = prepare_image_for_report(img_path)

                cell_img = tbl.rows[1].cells[0]
                paragraph = cell_img.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                paragraph.add_run().add_picture(
                    str(img_processed),
                    width=Cm(16)  
                )

            except Exception as e:
                print("Erro SETUP:", e)

            #  espaço entre tabelas (em vez de page break)
            last_paragraph = doc.add_paragraph("")


        # remover tabela base
        try:
            doc._body._body.remove(template_table._tbl)
        except Exception:
            pass

        nome_clean = teste.teste.teste.upper().replace(" ", "_")
        output_path = output_dir / f"{ensaio_numero}_{numero_pasta}_{nome_clean}_Setup.docx"

        doc.save(output_path)

    # =================================================
    # RESULTADO
    # =================================================
    return jsonify({
        "success": True,
        "message_key": "msg.relatorios_criados_sucesso",
        "params": {"path": str(output_dir)}
    })



@relatorios_bp.route('/relatorios/verificar', methods=['POST'])
def verificar_relatorio():

    data = request.get_json()

    ensaio_numero = data.get("ensaio")
    testes_sel = data.get("testes", [])

    if not ensaio_numero:
        return jsonify({"error_key": "msg.preencha_ensaio"}), 400

    ensaio = Ensaio.query.filter_by(ensaio=ensaio_numero).first()
    if not ensaio:
        return jsonify({"error_key": "msg.ensaio_nao_encontrado"}), 404

    # =================================================
    # BASE PATH
    # =================================================
    base_path = resolver_base_path(ensaio)

    if not base_path:
        return jsonify({
            "ok": False,
            "erros": ["Pasta do ensaio não encontrada"]
        })

    erros = []

    # =================================================
    # BEFORE TEST
    # =================================================
    before_test_base = base_path / "BEFORE TEST"

    if not before_test_base.exists():
        erros.append("Pasta BEFORE TEST não existe")
        return jsonify({"ok": False, "erros": erros})

    pecas = obter_pecas(before_test_base)

    if not pecas:
        erros.append("Não existem peças no BEFORE TEST")
        return jsonify({"ok": False, "erros": erros})

    # =================================================
    # TESTES
    # =================================================
    testes = Testes.query.filter_by(ensaio_id=ensaio.id).all()

    # lista base (todos os que criam pasta)
    testes_com_pasta = [
        t for t in sorted(testes, key=lambda x: x.ordem or 0)
        if t.teste and t.teste.criarpasta
    ]

    # lista final (filtrada pelo utilizador)
    testes_validos = [
        t for t in testes_com_pasta
        if not testes_sel or t.ordem in testes_sel
    ]

    if not testes_validos:
        erros.append("Nenhum teste selecionado válido")
        return jsonify({"ok": False, "erros": erros})

    previous_after_test_path = None

    # =================================================
    # VERIFICAÇÃO POR TESTE
    # =================================================
    for teste in testes_validos:

        numero_pasta = testes_com_pasta.index(teste) + 1

        nome_pasta = f"{numero_pasta} - {teste.teste.teste.upper()}"

        pasta_teste = find_child_dir(base_path, nome_pasta)


        if not pasta_teste:
            erros.append(f"Pasta do teste não encontrada: {nome_pasta}")
            continue

        after_base = pasta_teste / "AFTER TEST"

        if not after_base.exists():
            erros.append(f"AFTER TEST não existe para {nome_pasta}")

        # =================================================
        # VERIFICAR IMAGENS POR PEÇA
        # =================================================
        for peca in pecas:

            before_imgs = []
            after_imgs = []

            # BEFORE
            if previous_after_test_path is None:
                before_path = before_test_base / peca.name
            else:
                before_path = previous_after_test_path / peca.name

            if before_path.exists():
                before_imgs = list(before_path.glob("*.*"))

            if not before_imgs:
                erros.append(f"Sem imagens BEFORE: {peca.name} ({nome_pasta})")

            # AFTER
            after_piece_path = after_base / peca.name if after_base.exists() else None

            if after_piece_path and after_piece_path.exists():
                after_imgs = list(after_piece_path.glob("*.*"))

            if not after_imgs:
                erros.append(f"Sem imagens AFTER: {peca.name} ({nome_pasta})")

        previous_after_test_path = after_base if after_base.exists() else None

    # =================================================
    # RESULTADO FINAL
    # =================================================
    return jsonify({
        "ok": len(erros) == 0,
        "erros": erros
    })

